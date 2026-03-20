"""
Step 2 — PDF Parser + Text Cleaner
Handles: normal PDFs, badly formatted PDFs, scanned PDFs, edge cases
"""

import re
import os
from enum import Enum
import pdfplumber
import fitz  # PyMuPDF
from docx import Document as DocxDocument

class ParseStatus(Enum):
    OK = "ok"
    EMPTY = "parsing_failed_empty"
    SCANNED = "parsing_failed_scanned"
    ERROR = "parsing_failed_error"


def clean_text(text: str) -> str:
    """Remove noise from extracted text."""
    # Remove excessive whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]{2,}', ' ', text)
    # Remove page numbers (standalone numbers on a line)
    text = re.sub(r'^\s*\d+\s*$', '', text, flags=re.MULTILINE)
    # Remove URLs
    text = re.sub(r'http\S+', '', text)
    # Remove special characters but keep useful punctuation
    text = re.sub(r'[^\w\s\n\.\,\-\+\#\/\@\(\)]', ' ', text)
    return text.strip()


def extract_with_pdfplumber(path: str) -> str:
    """Primary extraction method."""
    text = ""
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text


def extract_with_pymupdf(path: str) -> str:
    """Fallback extraction method."""
    text = ""
    doc = fitz.open(path)
    for page in doc:
        text += page.get_text() + "\n"
    doc.close()
    return text

def extract_with_docx(path: str) -> str:
    """Extract text from .docx files."""
    doc = DocxDocument(path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " "
            text += "\n"
    return text


def detect_sections(text: str) -> dict:
    """Split resume text into sections."""
    sections = {
        "skills": "",
        "experience": "",
        "education": "",
        "summary": "",
        "projects": "",
    }

    # Section header patterns
    patterns = {
        "skills": r"(skills?|technical skills?|core competencies|technologies)",
        "experience": r"(experience|work experience|employment|professional experience|internship)",
        "education": r"(education|academic|qualification|degree)",
        "summary": r"(summary|objective|profile|about me|overview)",
        "projects": r"(projects?|personal projects?|key projects?)",
    }

    lines = text.split('\n')
    current_section = "summary"
    section_content = {k: [] for k in sections}

    for line in lines:
        line_lower = line.lower().strip()
        matched = False
        for section, pattern in patterns.items():
            if re.search(pattern, line_lower) and len(line_lower) < 50:
                current_section = section
                matched = True
                break
        if not matched:
            section_content[current_section].append(line)

    for section in sections:
        sections[section] = '\n'.join(section_content[section]).strip()

    return sections


def looks_scanned(text: str) -> bool:
    """Detect if PDF is image-based (scanned)."""
    if not text:
        return True
    # Very little text extracted = likely scanned
    words = text.split()
    return len(words) < 30


def parse_pdf(path: str) -> dict:
    """
    Main parse function. Handles PDF and DOCX files.
    Returns structured dict with status flag.
    Always returns a dict — never crashes.
    """
    filename = os.path.basename(path)
    ext = os.path.splitext(filename)[1].lower()

    try:
        # Route to correct extractor based on file type
        if ext == ".docx":
            text = extract_with_docx(path)
        else:
            # Try pdfplumber first
            text = extract_with_pdfplumber(path)
            # Fallback to PyMuPDF if pdfplumber gives too little
            if not text or len(text.strip()) < 100:
                text = extract_with_pymupdf(path)

        # Check if scanned
        if looks_scanned(text):
            return {
                "status": ParseStatus.SCANNED,
                "filename": filename,
                "raw": "",
                "sections": {},
                "error": "PDF appears to be scanned/image-based. Text extraction failed."
            }

        # Clean the text
        cleaned = clean_text(text)

        # Check if still too short after cleaning
        if len(cleaned.strip()) < 50:
            return {
                "status": ParseStatus.EMPTY,
                "filename": filename,
                "raw": "",
                "sections": {},
                "error": "Extracted text is too short. PDF may be corrupted."
            }

        # Detect sections
        sections = detect_sections(cleaned)

        return {
            "status": ParseStatus.OK,
            "filename": filename,
            "raw": cleaned,
            "sections": sections,
            "error": None
        }

    except Exception as e:
        return {
            "status": ParseStatus.ERROR,
            "filename": filename,
            "raw": "",
            "sections": {},
            "error": str(e)
        }

def parse_all_resumes(resumes_dir: str) -> list:
    """Parse all PDFs in a directory."""
    results = []
    pdf_files = [f for f in os.listdir(resumes_dir)
                 if f.endswith('.pdf') or f.endswith('.docx')]

    if not pdf_files:
        print("No PDF files found in", resumes_dir)
        return results

    print(f"Found {len(pdf_files)} resume(s). Parsing...\n")

    for filename in pdf_files:
        path = os.path.join(resumes_dir, filename)
        result = parse_pdf(path)
        results.append(result)

        status = result["status"].value
        if result["status"] == ParseStatus.OK:
            word_count = len(result["raw"].split())
            print(f"  OK      {filename} ({word_count} words extracted)")
        else:
            print(f"  FAILED  {filename} — {result['error']}")

    return results


# ── Quick test ──────────────────────────────────────────────
if __name__ == "__main__":
    import sys

    resumes_dir = "data/resumes"
    results = parse_all_resumes(resumes_dir)

    print(f"\n{'='*50}")
    print(f"Parsed {len(results)} resumes")
    ok = sum(1 for r in results if r["status"] == ParseStatus.OK)
    failed = len(results) - ok
    print(f"  Successful : {ok}")
    print(f"  Failed     : {failed}")

    if results:
        print(f"\n--- Sample sections from first OK resume ---")
        for r in results:
            if r["status"] == ParseStatus.OK:
                print(f"\nFile: {r['filename']}")
                for section, content in r["sections"].items():
                    preview = content[:100].replace('\n', ' ')
                    print(f"  [{section}]: {preview}...")
                break